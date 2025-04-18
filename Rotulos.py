from docx import Document
from docx2pdf import convert
from datetime import datetime, timedelta
import os
import time

def substituir_texto(documento, antigo, novo):
    for paragrafo in documento.paragraphs:
        if antigo in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(antigo, novo)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if antigo in celula.text:
                    celula.text = celula.text.replace(antigo, novo)

def calcular_dia_juliano(data):
    return data.timetuple().tm_yday

def gerar_codigo_lote(data):
    ano = data.strftime("%y")
    dia_juliano = f"{calcular_dia_juliano(data):03d}"
    return f"L{ano}{dia_juliano}076YY"

def gerar_pdfs_em_serie(caminho_arquivo_word, pasta_saida):
    data_inicial = datetime(2025, 1, 1)

    if not os.path.exists(pasta_saida):
        os.makedirs(pasta_saida)

    for dia in range(365):
        try:
            data_atual = data_inicial + timedelta(days=dia)
            data_fabricacao = data_atual.strftime("%d/%m/%Y")
            data_validade = (data_atual + timedelta(days=60)).strftime("%d/%m/%Y")
            codigo_lote = gerar_codigo_lote(data_atual)

            documento = Document(caminho_arquivo_word)

            substituir_texto(documento, "DATA_DE_FABRICAÇÃO", data_fabricacao)
            substituir_texto(documento, "DATA_DE_VALIDADE", data_validade)
            substituir_texto(documento, "LOTE", codigo_lote)

            nome_arquivo_word = f"Rótulo_Pó_{data_atual.strftime('%Y-%m-%d')}.docx"
            caminho_arquivo_word_modificado = os.path.join(pasta_saida, nome_arquivo_word)
            documento.save(caminho_arquivo_word_modificado)

            nome_arquivo_pdf = f"Rótulo_Pó_{data_atual.strftime('%Y-%m-%d')}.pdf"
            caminho_arquivo_pdf = os.path.join(pasta_saida, nome_arquivo_pdf)
            

            try:
                convert(caminho_arquivo_word_modificado, caminho_arquivo_pdf)
                print(f"Arquivo gerado com sucesso: {nome_arquivo_pdf}")
            except Exception as e:
                print(f"Erro ao converter para PDF: {e}")
                os.system('taskkill /f /im winword.exe')
                time.sleep(2)
            
        except Exception as e:
            print(f"Erro ao processar o dia {dia}: {e}")
            continue
caminho_arquivo_word = r"C:\Users\jekfi\OneDrive\Documentos\Rotulo\Rótulo de pó WW.docx"
pasta_saida = r"C:\Users\jekfi\OneDrive\Documentos\Teste arquivos"
gerar_pdfs_em_serie(caminho_arquivo_word, pasta_saida)